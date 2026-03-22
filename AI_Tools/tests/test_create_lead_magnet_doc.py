"""
Tests for create_lead_magnet_doc.py
Run: py -3 24_Tools/tests/test_create_lead_magnet_doc.py
"""
import sys
import tempfile
from pathlib import Path

# Allow import from 24_Tools/
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import create_lead_magnet_doc as cmd
from docx import Document


def test_new_doc_has_correct_margins():
    doc = cmd._new_doc()
    s = doc.sections[0]
    from docx.shared import Inches
    assert s.left_margin == Inches(1.0), "Left margin should be 1 inch"
    assert s.right_margin == Inches(1.0), "Right margin should be 1 inch"
    print("  [OK] margins ok")


def test_all_guides_render_without_error():
    for key in cmd.GUIDES:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            original_out = cmd.OUT_DIR
            cmd.OUT_DIR = tmp_path
            try:
                cmd.render(key)
                out_file = tmp_path / cmd.GUIDES[key][0]
                assert out_file.exists(), f"{key}: output file not created"
                # Verify it's a valid docx
                doc = Document(str(out_file))
                assert len(doc.paragraphs) > 0 or len(doc.tables) > 0, f"{key}: document appears empty"
                print(f"  [OK] {key}: rendered ok ({out_file.stat().st_size // 1024} KB)")
            finally:
                cmd.OUT_DIR = original_out


def test_components_do_not_crash():
    doc = cmd._new_doc()
    cmd.add_cover(doc, "Test Title\nLine Two", "Test subtitle text.")
    cmd.add_intro_box(doc, "Intro box text here.")
    cmd.add_section_label(doc, "Section Label")
    cmd.add_heading3(doc, "Heading Three")
    cmd.add_body(doc, "Body text with **bold** marker.")
    cmd.add_bullets(doc, ["Item 1", "**Bold** item 2", "Item 3"])
    cmd.add_numbered_list(doc, ["First item *italic*", "Second item"])
    cmd.add_step(doc, 1, "Step Title", "Step body text.")
    cmd.add_code_block(doc, "code here", label="Code Label", style="default")
    cmd.add_code_block(doc, "user prompt", style="user")
    cmd.add_code_block(doc, "ai output", style="ai")
    cmd.add_code_block(doc, "system block", style="system")
    cmd.add_formula_box(doc, "[formula] + [here]")
    cmd.add_table(doc,
        headers=["Col A", "Col B"],
        rows=[["val 1", "val 2"], ["val **bold**", "val 4"]])
    cmd.add_scenario(doc, 1, "Scenario Title", "Scenario desc.",
                     "prompt text", "ai output text")
    cmd.add_flow_list(doc, ["Step one", "Step two", "**Final step**"])
    cmd.add_callout(doc, "Warning text here.", prefix="Note:")
    cmd.add_divider(doc)
    cmd.add_cta_box(doc, "CTA Heading", "CTA body text.", "DM Keyword")
    cmd.add_footer(doc)
    print("  [OK] all components render without error")


if __name__ == "__main__":
    print("\nRunning create_lead_magnet_doc tests...\n")
    tests = [
        test_new_doc_has_correct_margins,
        test_components_do_not_crash,
        test_all_guides_render_without_error,
    ]
    failed = 0
    for t in tests:
        try:
            print(f"[{t.__name__}]")
            t()
        except Exception as e:
            print(f"  [FAIL] FAILED: {e}")
            failed += 1
    print(f"\n{'All tests passed.' if not failed else f'{failed} test(s) failed.'}")
    sys.exit(failed)
